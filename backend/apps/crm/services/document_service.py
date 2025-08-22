# ============================================================================
# backend/apps/crm/services/document_service.py - Advanced Document Management Service
# ============================================================================

import os
import mimetypes
import hashlib
import magic
from typing import Dict, List, Any, Optional, Tuple, Union
from datetime import datetime, timedelta
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.db import transaction
from django.utils import timezone
from django.conf import settings
from django.core.exceptions import ValidationError
from PIL import Image
import pypdf
import docx
import openpyxl
from io import BytesIO
import logging

from .base import BaseService, ServiceException
from ..models import (
    Document, DocumentCategory, DocumentShare, DocumentVersion,
    AuditTrail
)

logger = logging.getLogger(__name__)


class DocumentProcessingError(ServiceException):
    """Document processing specific errors"""
    pass


class DocumentSecurityError(ServiceException):
    """Document security specific errors"""
    pass


class DocumentService(BaseService):
    """Advanced document management service with AI capabilities"""
    
    # File type configurations
    ALLOWED_EXTENSIONS = {
        'documents': ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt'],
        'spreadsheets': ['.xls', '.xlsx', '.csv', '.ods'],
        'images': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
        'presentations': ['.ppt', '.pptx', '.odp'],
        'archives': ['.zip', '.rar', '.7z', '.tar', '.gz'],
        'other': ['.xml', '.json', '.yaml', '.yml']
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    THUMBNAIL_SIZE = (300, 300)
    VIRUS_SCAN_ENABLED = getattr(settings, 'DOCUMENT_VIRUS_SCAN', False)
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.ai_categorizer = AIDocumentCategorizer()
        self.content_extractor = DocumentContentExtractor()
        self.security_scanner = DocumentSecurityScanner()
    
    # ============================================================================
    # CORE DOCUMENT OPERATIONS
    # ============================================================================
    
    @transaction.atomic
    def upload_document(self, file, metadata: Dict = None, 
                       category_id: int = None, tags: List[str] = None,
                       security_level: str = 'normal') -> Document:
        """
        Upload document with comprehensive processing and validation
        
        Args:
            file: Django uploade metadata
            category_id: Document category
            tags: List of tags
            security_level: Security level (low, normal, high, confidential)
        
        Returns:
            Document instance
        """
        self.context.operation = 'upload_document'
        
        try:
            # Validate file
            self._validate_file(file, security_level)
            
            # Extract metadata
            file_metadata = self._extract_file_metadata(file)
            
            # Generate unique filename
            unique_filename = self._generate_unique_filename(file.name)
            
            # Store file
            file_path = self._store_file(file, unique_filename)
            
            # Extract content for indexing
            extracted_content = self.content_extractor.extract_content(file)
            
            # AI categorization if no category provided
            suggested_category = None
            if not category_id and extracted_content:
                suggested_category = self.ai_categorizer.suggest_category(
                    filename=file.name,
                    content=extracted_content.get('text', ''),
                    tenant=self.tenant
                )
            
            # Create document record
            document_data = {
                'tenant': self.tenant,
                'name': metadata.get('name', os.path.splitext(file.name)[0]),
                'description': metadata.get('description', ''),
                'file_name': file.name,
                'file_path': file_path,
                'file_size': file.size,
                'mime_type': file_metadata['mime_type'],
                'file_hash': file_metadata['hash'],
                'category_id': category_id or (suggested_category.id if suggested_category else None),
                'uploaded_by': self.user,
                'security_level': security_level,
                'is_active': True,
                'metadata': {
                    **file_metadata,
                    **(metadata or {}),
                    'extracted_content': extracted_content,
                    'ai_suggestions': {
                        'suggested_category': suggested_category.name if suggested_category else None
                    }
                }
            }
            
            document = Document.objects.create(**document_data)
            
            # Create initial version
            self._create_document_version(document, file_path, 'Initial upload')
            
            # Add tags
            if tags:
                self._add_tags_to_document(document, tags)
            
            # Generate thumbnail if image
            if file_metadata['is_image']:
                self._generate_thumbnail(document)
            
            # Index for search
            self._index_document_for_search(document, extracted_content)
            
            # Log activity
            self.log_activity(
                'document_uploaded',
                'Document',
                document.id,
                {
                    'file_name': file.name,
                    'file_size': file.size,
                    'security_level': security_level,
                    'category': suggested_category.name if suggested_category else None
                }
            )
            
            return document
            
        except Exception as e:
            logger.error(f"Document upload failed: {e}", exc_info=True)
            raise DocumentProcessingError(f"Upload failed: {str(e)}")
    
    def download_document(self, document_id: int, version_number: int = None) -> Dict:
        """
        Download document with access control and audit logging
        
        Args:
            document_id: Document ID
            version_number: Specific version (latest if None)
        
        Returns:
            Dict with file info and download URL
        """
        try:
            document = Document.objects.select_related('category', 'uploaded_by').get(
                id=document_id,
                tenant=self.tenant
            )
            
            # Check access permissions
            self._check_document_access(document, 'read')
            
            # Get specific version or latest
            if version_number:
                version = document.versions.filter(version_number=version_number).first()
                if not version:
                    raise DocumentProcessingError(f"Version {version_number} not found")
                file_path = version.file_path
            else:
                file_path = document.file_path
                version = document.versions.order_by('-version_number').first()
            
            # Generate secure download URL
            download_url = self._generate_secure_download_url(document, file_path)
            
            # Log download activity
            self.log_activity(
                'document_downloaded',
                'Document',
                document.id,
                {
                    'version_number': version.version_number if version else None,
                    'file_name': document.file_name,
                    'download_method': 'secure_url'
                }
            )
            
            # Update access statistics
            self._update_document_stats(document, 'download')
            
            return {
                'document_id': document.id,
                'name': document.name,
                'file_name': document.file_name,
                'file_size': document.file_size,
                'mime_type': document.mime_type,
                'download_url': download_url,
                'version_number': version.version_number if version else 1,
                'expires_at': timezone.now() + timedelta(hours=1)  # URL expires in 1 hour
            }
            
        except Document.DoesNotExist:
            raise DocumentProcessingError("Document not found")
    
    @transaction.atomic
    def update_document_version(self, document_id: int, new_file, 
                               version_notes: str = None) -> DocumentVersion:
        """
        Update document with new version while preserving history
        
        Args:
            document_id: Document ID
            new_file: New file upload
            version_notes: Version change notes
        
        Returns:
            New DocumentVersion instance
        """
        try:
            document = Document.objects.get(id=document_id, tenant=self.tenant)
            
            # Check update permissions
            self._check_document_access(document, 'update')
            
            # Validate new file
            self._validate_file(new_file, document.security_level)
            
            # Get current version number
            current_version = document.versions.order_by('-version_number').first()
            new_version_number = (current_version.version_number if current_version else 0) + 1
            
            # Extract new file metadata
            file_metadata = self._extract_file_metadata(new_file)
            
            # Store new version file
            version_filename = f"{document.id}_v{new_version_number}_{new_file.name}"
            new_file_path = self._store_file(new_file, version_filename)
            
            # Extract content from new version
            extracted_content = self.content_extractor.extract_content(new_file)
            
            # Update document main record
            document.file_name = new_file.name
            document.file_path = new_file_path
            document.file_size = new_file.size
            document.mime_type = file_metadata['mime_type']
            document.file_hash = file_metadata['hash']
            document.updated_by = self.user
            document.updated_at = timezone.now()
            
            # Update metadata
            if .update({
                    'extracted_content': extracted_content,
                    'version_history': document.metadata.get('version_history', [])
                })
            else:
                document.metadata = {'extracted_content': extracted_content}
            
            document.save()
            
            # Create new version record
            version = DocumentVersion.objects.create(
                document=document,
                version_number=new_version_number,
                file_path=new_file_path,
                file_size=new_file.size,
                file_hash=file_metadata['hash'],
                version_notes=version_notes or f"Version {new_version_number}",
                created_by=self.user,
                tenant=self.tenant
            )
            
            # Update search index
            self._update_document_search_index(document, extracted_content)
            
            # Generate new thumbnail if image
            if file_metadata['is_image']:
                self._generate_thumbnail(document)
            
            # Log version update
            self.log_activity(
                'document_version_updated',
                'Document',
                document.id,
                {
                    'new_version': new_version_number,
                    'file_name': new_file.name,
                    'version_notes': version_notes
                }
            )
            
            return version
            
        except Document.DoesNotExist:
            raise DocumentProcessingError("Document not found")
    
    # ============================================================================
    # DOCUMENT SHARING AND COLLABORATION
    # ============================================================================
    
    @transaction.atomic
    def share_document(self, document_id: int, share_with: List[Dict], 
                      permissions: List[str], expires_at: datetime = None,
                      message: str = None) -> List[DocumentShare]:
        """
        Share document with specific users/groups with granular permissions
        
        Args:
            document_id: Document ID
            share_with: List of {'type': 'user'/'group', 'id': int}
            permissions: List of permissions ['read', 'write', 'delete', 'share']
            expires_at: Share expiration
            message: Share message
        
        Returns:
            List of DocumentShare instances
        """
        try:
            document = Document.objects.get(id=document_id, tenant=self.tenant)
            
            # Check sharing permissions
            self._check_document_access(document, 'share')
            
            shares_created = []
            
            for share_target in share_with:
                share_type = share_target['type']
                target_id = share_target['id']
                
                # Validate share target
                if share_type == 'user':
                    # Ensure user belongs to same tenant
                    from apps.auth.models import Membership
                    if not Membership.objects.filter(
                        user_id=target_id, 
                        tenant=self.tenant
                    ).exists():
                        continue
                
                # Create or update share
                share, created = DocumentShare.objects.update_or_create(
                    document=document,
                    shared_with_type=share_type,
                    shared_with_id=target_id,
                    defaults={
                        'shared_by': self.user,
                        'permissions': permissions,
                        'expires_at': expires_at,
                        'message': message,
                        'is_active': True,
                        'tenant': self.tenant
                    }
                )
                
                shares_created.append(share)
                
                # Send notification about share
                if created:
                    self._send_share_notification(share, message)
            
            # Log sharing activity
            self.log_activity(
                'document_shared',
                'Document',
                document.id,
                {
                    'shared_with_count': len(shares_created),
                    'permissions': permissions,
                    'expires_at': expires_at.isoformat() if expires_at else None
                }
            )
            
            return shares_created
            
        except Document.DoesNotExist:
            raise DocumentProcessingError("Document not found")
    
    def get_shared_documents(self, user_id: int = None) -> List[Dict]:
        """
        Get documents shared with specific user or current user
        
        Args:
            user_id: User ID (current user if None)
        
        Returns:
            List of shared document information
        """
        target_user_id = user_id or self.user.id
        
        # Get active shares for user
        shares = DocumentShare.objects.select_related(
            'document', 'shared_by', 'document__category'
        ).filter(
            shared_with_type='user',
            shared_with_id=target_user_id,
            is_active=True,
            tenant=self.tenant
        ).filter(
            # Check expiration
            models.Q(expires_at__isnull=True) | models.Q(expires_at__gt=timezone.now())
        ).order_by('-created_at')
        
        shared_docs = []
        for share in shares:
            doc = share.document
            shared_docs.append({
                'share_id': share.id,
                'document_id': doc.id,
                'name': doc.name,
                'file_name': doc.file_name,
                'file_size': doc.file_size,
                'mime_type': doc.mime_type,
                'category': doc.category.name if doc.category else None,
                'shared_by': share.shared_by.get_full_name(),
                'shared_at': share.created_at,
                'expires_at': share.expires_at,
                'permissions': share.permissions,
                'message': share.message,
                'thumbnail_url': self._get_thumbnail_url(doc) if doc.metadata.get('has_thumbnail') else None
            })
        
        return shared_docs
    
    # ============================================================================
    # DOCUMENT ORGANIZATION AND SEARCH
    # ============================================================================
    
    def create_category(self, name: str, description: str = None, 
                       parent_id: int = None, color: str = None,
                       auto_categorize_rules: Dict = None) -> DocumentCategory:
        """
        Create document category with AI auto-categorization rules
        
        Args:
            name: Category name
            description: Category description
            parent_id: Parent category ID
            color: Category color
            auto_categorize_rules: AI rules for auto-categorization
        
        Returns:
            DocumentCategory instance
        """
        try:
            self.validate_user_permission('crm.add_documentcategory')
            
            # Check for duplicate names at same level
            existing = DocumentCategory.objects.filter(
                name__iexact=name,
                parent_id=parent_id,
                tenant=self.tenant
            ).exists()
            
            if existing:
                raise ValidationError(f"Category '{name}' already exists at this level")
            
            category = DocumentCategory.objects.create(
                name=name,
                description=description or '',
                parent_id=parent_id,
                color=color or '#6366f1',
                auto_categorize_rules=auto_categorize_rules or {},
                tenant=self.tenant,
                created_by=self.user
            )
            
            # Update AI categorizer with new category
            if auto_categorize_rules:
                self.ai_categorizer.update_category_rules(category)
            
            self.log_activity(
                'document_category_created',
                'DocumentCategory',
                category.id,
                {'name': name, 'parent_id': parent_id}
            )
            
            return category
            
        except Exception as e:
            raise DocumentProcessingError(f"Category creation failed: {str(e)}")
    
    def smart_search_documents(self, query: str, filters: Dict = None,
                             limit: int = 50, include_content: bool = True) -> Dict:
        """
        Advanced document search with content analysis and AI ranking
        
        Args:
            query: Search query
            filters: Additional filters
            limit: Result limit
            include_content: Whether to search in content
        
        Returns:
            Dict with results and metadata
        """
        try:
            from django.contrib.postgres.search import SearchVector, SearchQuery, SearchRank
            from django.db.models import Q, F
            
            # Build base queryset
            queryset = Document.objects.select_related(
                'category', 'uploaded_by'
            ).filter(
                tenant=self.tenant,
                is_active=True
            )
            
            # Apply filters
            if filters:
                if 'category_id' in filters:
                    queryset = queryset.filter(category_id=filters['category_id'])
                
                if 'file_type' in filters:
                    file_extensions = self.ALLOWED_EXTENSIONS.get(filters['file_type'], [])
                    extension_q = Q()
                    for ext in file_extensions:
                        extension_q |= Q(file_name__iendswith=ext)
                    queryset = queryset.filter(extension_q)
                
                if 'date_range' in filters:
                    date_range = filters['date_range']
                    queryset = queryset.filter(created_at__range=date_range)
                
                if 'security_level' in filters:
                    queryset = queryset.filter(security_level=filters['security_level'])
            
            # Text search
            search_results = []
            
            if query.strip():
                # Search in document fields
                search_vector = SearchVector('name', weight='A') + \
                              SearchVector('description', weight='B') + \
                              SearchVector('file_name', weight='C')
                
                if include_content:
                    # Search in extracted content (stored in metadata)
                    content_query = Q(metadata__extracted_content__text__icontains=query)
                    queryset = queryset.filter(content_query)
                
                search_query = SearchQuery(query)
                
                search_results = queryset.annotate(
                    search=search_vector,
                    rank=SearchRank(search_vector, search_query)
                ).filter(
                    Q(search=search_query) | Q(name__icontains=query) | 
                    Q(description__icontains=query) | Q(file_name__icontains=query)
                ).order_by('-rank', '-created_at')[:limit]
            else:
                search_results = queryset.order_by('-created_at')[:limit]
            
            # Format results
            results = []
            for doc in search_results:
                result = {
                    'id': doc.id,
                    'name': doc.name,
                    'description': doc.description,
                    'file_name': doc.file_name,
                    'file_size': doc.file_size,
                    'mime_type': doc.mime_type,
                    'category': {
                        'id': doc.category.id,
                        'name': doc.category.name,
                        'color': doc.category.color
                    } if doc.category else None,
                    'uploaded_by': doc.uploaded_by.get_full_name(),
                    'uploaded_at': doc.created_at,
                    'security_level': doc.security_level,
                    'thumbnail_url': self._get_thumbnail_url(doc) if doc.metadata.get('has_thumbnail') else None,
                    'relevance_score': getattr(doc, 'rank', 0)
                }
                
                # Add content snippets for search matches
                if include_content and query.strip():
                    content_snippet = self._extract_content_snippet(doc, query)
                    result['content_snippet'] = content_snippet
                
                results.append(result)
            
            # Search metadata
            search_metadata = {
                'query': query,
                'total_results': len(results),
                'search_time': time.time() - start_time if 'start_time' in locals() else 0,
                'filters_applied': filters or {},
                'suggestions': self._get_search_suggestions(query, results)
            }
            
            return {
                'results': results,
                'metadata': search_metadata
            }
            
        except Exception as e:
            logger.error(f"Document search failed: {e}", exc_info=True)
            raise DocumentProcessingError(f"Search failed: {str(e)}")
    
    # ============================================================================
    # DOCUMENT ANALYTICS AND INSIGHTS
    # ============================================================================
    
    def get_document_analytics(self, period: str = '30d') -> Dict:
        """
        Get comprehensive document analytics and insights
        
        Args:
            period: Analysis period ('7d', '30d', '90d', '1y')
        
        Returns:
            Dict with analytics data
        """
        try:
            # Calculate date range
            period_days = {'7d': 7, '30d': 30, '90d': 90, '1y': 365}
            days = period_days.get(period, 30)
            start_date = timezone.now() - timedelta(days=days)
            
            # Base queryset
            docs = Document.objects.filter(tenant=self.tenant, is_active=True)
            period_docs = docs.filter(created_at__gte=start_date)
            
            # Basic statistics
            total_documents = docs.count()
            period_uploads = period_docs.count()
            total_size = docs.aggregate(models.Sum('file_size'))['file_size__sum'] or 0
            
            # Document type distribution
            type_distribution = {}
            for doc in docs.select_related():
                mime_type = doc.mime_type.split('/')[0] if doc.mime_type else 'unknown'
                type_distribution[mime_type] = type_distribution.get(mime_type, 0) + 1
            
            # Category distribution
            category_stats = docs.values('category__name').annotate(
                count=models.Count('id'),
                total_size=models.Sum('file_size')
            ).order_by('-count')
            
            # Upload trends (daily for period)
            upload_trends = []
            for i in range(days):
                date = start_date + timedelta(days=i)
                day_uploads = period_docs.filter(created_at__date=date.date()).count()
                upload_trends.append({
                    'date': date.date().isoformat(),
                    'uploads': day_uploads
                })
            
            # Most active users
            active_users = docs.values('uploaded_by__first_name', 'uploaded_by__last_name').annotate(
                upload_count=models.Count('id'),
                total_size=models.Sum('file_size')
            ).order_by('-upload_count')[:10]
            
            # Security level distribution
            security_stats = docs.values('security_level').annotate(
                count=models.Count('id')
            )
            
            # Storage insights
            storage_insights = {
                'total_size_mb': round(total_size / (1024 * 1024), 2),
                'average_file_size_mb': round((total_size / total_documents) / (1024 * 1024), 2) if total_documents else 0,
                'largest_files': list(docs.order_by('-file_size')[:5].values(
                    'name', 'file_name', 'file_size', 'uploaded_by__first_name'
                ))
            }
            
            # AI insights
            ai_insights = self._generate_ai_document_insights(docs, period_docs)
            
            return {
                'overview': {
                    'total_documents': total_documents,
                    'period_uploads': period_uploads,
                    'total_size_mb': storage_insights['total_size_mb'],
                    'growth_rate': round((period_uploads / total_documents * 100), 2) if total_documents else 0
                },
                'type_distribution': type_distribution,
                'category_stats': list(category_stats),
                'upload_trends': upload_trends,
                'active_users': list(active_users),
                'security_distribution': list(security_stats),
                'storage_insights': storage_insights,
                'ai_insights': ai_insights,
                'period': period,
                'generated_at': timezone.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Analytics generation failed: {e}", exc_info=True)
            raise DocumentProcessingError(f"Analytics failed: {str(e)}")
    
    # ============================================================================
    # HELPER METHODS
    # ============================================================================
    
    def _validate_file(self, file, security_level: str = 'normal'):
        """Comprehensive file validation"""
        # Size validation
        if file.size > self.MAX_FILE_SIZE:
            raise DocumentSecurityError(f"File size exceeds limit ({self.MAX_FILE_SIZE} bytes)")
        
        # Extension validation
        file_ext = os.path.splitext(file.name)[1].lower()
        allowed_extensions = []
        for ext_list in self.ALLOWED_EXTENSIONS.values():
            allowed_extensions.extend(ext_list)
        
        if file_ext not in allowed_extensions:
            raise DocumentSecurityError(f"File type '{file_ext}' not allowed")
        
        # MIME type validation
        mime_type = magic.from_buffer(file.read(1024), mime=True)
        file.seek(0)  # Reset file pointer
        
        # Security scanning
        if security_level in ['high', 'confidential']:
            self._perform_security_scan(file)
    
    def _extract_file_metadata(self, file) -> Dict:
        """Extract comprehensive file metadata"""
        file_ext = os.path.splitext(file.name)[1].lower()
        mime_type = magic.from_buffer(file.read(1024), mime=True)
        file.seek(0)
        
        # Calculate hash
        hash_md5 = hashlib.md5()
        for chunk in iter(lambda: file.read(4096), b""):
            hash_md5.update(chunk)
        file.seek(0)
        
        metadata = {
            'extension': file_ext,
            'mime_type': mime_type,
            'hash': hash_md5.hexdigest(),
            'is_image': mime_type.startswith('image/'),
            'is_document': mime_type in ['application/pdf', 'application/msword', 
                                       'application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
            'uploaded_at': timezone.now().isoformat()
        }
        
        return metadata
    
    def _generate_unique_filename(self, original_name: str) -> str:
        """Generate unique filename to prevent conflicts"""
        name, ext = os.path.splitext(original_name)
        timestamp = int(timezone.now().timestamp())
        unique_id = str(uuid.uuid4())[:8]
        return f"{name}_{timestamp}_{unique_id}{ext}"
    
    def _store_file(self, file, filename: str) -> str:
        """Store file using Django's storage backend"""
        file_path = f"documents/{self.tenant.id}/{timezone.now().year}/{timezone.now().month}/{filename}"
        return default_storage.save(file_path, file)
    
    def _check_document_access(self, document, operation: str):
        """Check document access permissions"""
        self.validate_tenant_access(document)
        
        # Owner always has access
        if document.uploaded_by == self.user:
            return True
        
        # Check shared access
        if operation in ['read', 'download']:
            share = DocumentShare.objects.filter(
                document=document,
                shared_with_type='user',
                shared_with_id=self.user.id,
                is_active=True
            ).filter(
                models.Q(expires_at__isnull=True) | models.Q(expires_at__gt=timezone.now())
            ).first()
            
            if share and operation in share.permissions:
                return True
        
        # Check general permissions
        required_permission = f'crm.{operation}_document'
        self.validate_user_permission(required_permission, document)
    
    def _create_document_version(self, document, file_path: str, notes: str):
        """Create initial document version"""
        return DocumentVersion.objects.create(
            document=document,
            version_number=1,
            file_path=file_path,
            file_size=document.file_size,
            file_hash=document.file_hash,
            version_notes=notes,
            created_by=self.user,
            tenant=self.tenant
        )
    
    def _generate_ai_document_insights(self, all_docs, period_docs) -> Dict:
        """Generate AI-powered insights about documents"""
        insights = {}
        
        try:
            # Trend analysis
            if period_docs.count() > 0:
                avg_daily_uploads = period_docs.count() / 30  # Assuming 30-day period
                if avg_daily_uploads > 5:
                    insights['upload_trend'] = 'High activity - consider storage optimization'
                elif avg_daily_uploads < 1:
                    insights['upload_trend'] = 'Low activity - encourage document sharing'
                else:
                    insights['upload_trend'] = 'Normal activity levels'
            
            # Category recommendations
            uncategorized_count = all_docs.filter(category__isnull=True).count()
            if uncategorized_count > all_docs.count() * 0.3:  # More than 30% uncategorized
                insights['organization'] = f'{uncategorized_count} documents need categorization'
            
            # Security recommendations
            low_security_sensitive = all_docs.filter(
                security_level='normal',
                file_name__iregex=r'.*(confidential|private|secret|password).*'
            ).count()
            
            if low_security_sensitive > 0:
                insights['security'] = f'{low_security_sensitive} documents may need higher security levels'
            
        except Exception as e:
            logger.error(f"AI insights generation failed: {e}")
            insights['error'] = 'AI insights temporarily unavailable'
        
        return insights


class AIDocumentCategorizer:
    """AI-powered document categorization"""
    
    def suggest_category(self, filename: str, content: str, tenant) -> Optional:
        """Suggest document category based on filename and content"""
        try:
            # Simple rule-based categorization (can be enhanced with ML)
            filename_lower = filename.lower()
            content_lower = content.lower() if content else ""
            
            categories = DocumentCategory.objects.filter(tenant=tenant)
            
            # Keywords mapping
            category_keywords = {
                'contract': ['contract', 'agreement', 'terms', 'conditions'],
                'invoice': ['invoice', 'bill', 'payment', 'amount due'],
                'report': ['report', 'analysis', 'summary', 'findings'],
                'presentation': ['presentation', 'slides', 'deck'],
                'legal': ['legal', 'law', 'court', 'litigation'],
                'hr': ['employee', 'hr', 'human resources', 'personnel'],
                'financial': ['financial', 'budget', 'accounting', 'expense']
            }
            
            for category in categories:
                category_name_lower = category.name.lower()
                
                # Check if category name or keywords match
                if category_name_lower in filename_lower or category_name_lower in content_lower:
                    return category
                
                # Check predefined keywords
                keywords = category_keywords.get(category_name_lower, [])
                for keyword in keywords:
                    if keyword in filename_lower or keyword in content_lower:
                        return category
            
            return None
            
        except Exception as e:
            logger.error(f"AI categorization failed: {e}")
            return None


class DocumentContentExtractor:
    """Extract text content from various document types"""
    
    def extract_content(self, file) -> Dict:
        """Extract content from uploaded file"""
        try:
            content = {'text': '', 'metadata': {}}
            mime_type = magic.from_buffer(file.read(1024), mime=True)
            file.seek(0)
            
            if mime_type == 'application/pdf':
                content = self._extract_pdf_content(file)
            elif mime_type in ['application/msword', 
                             'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                content = self._extract_word_content(file)
            elif mime_type in ['application/vnd.ms-excel',
                             'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
                content = self._extract_excel_content(file)
            elif mime_type == 'text/plain':
                content = self._extract_text_content(file)
            
            return content
            
        except Exception as e:
            logger.error(f"Content extraction failed: {e}")
            return {'text': '', 'metadata': {'extraction_error': str(e)}}
    
    def _extract_pdf_content(self, file) -> Dict:
        """Extract content from PDF"""
        try:
            pdf_reader = pypdf.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                'text': text.strip(),
                'metadata': {
                    'page_count': len(pdf_reader.pages),
                    'content_type': 'pdf'
                }
            }
        except Exception as e:
            return {'text': '', 'metadata': {'error': str(e)}}
    
    def _extract_word_content(self, file) -> Dict:
        """Extract content from Word documents"""
        try:
            doc = docx.Document(file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                'text': text.strip(),
                'metadata': {
                    'paragraph_count': len(doc.paragraphs),
                    'content_type': 'word'
                }
            }
        except Exception as e:
            return {'text': '', 'metadata': {'error': str(e)}}
    
    def _extract_excel_content(self, file) -> Dict:
        """Extract content from Excel files"""
        try:
            workbook = openpyxl.load_workbook(file)
            text = ""
            
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' '.join([str(cell) for cell in row if cell is not None])
                    text += row_text + "\n"
            
            return {
                'text': text.strip(),
                'metadata': {
                    'sheet_count': len(workbook.worksheets),
                    'content_type': 'excel'
                }
            }
        except Exception as e:
            return {'text': '', 'metadata': {'error': str(e)}}
    
    def _extract_text_content(self, file) -> Dict:
        """Extract content from text files"""
        try:
            text = file.read().decode('utf-8')
            return {
                'text': text,
                'metadata': {
                    'content_type': 'text',
                    'character_count': len(text)
                }
            }
        except Exception as e:
            return {'text': '', 'metadata': {'error': str(e)}}


class DocumentSecurityScanner:
    """Document security scanning functionality"""
    
    def scan_document(self, file) -> Dict:
        """Perform security scan on document"""
        results = {
            'safe': True,
            'threats_found': [],
            'recommendations': []
        }
        
        try:
            # Basic checks (can be enhanced with actual antivirus integration)
            filename = file.name.lower()
            
            # Check for suspicious extensions
            dangerous_extensions = ['.exe', '.bat', '.cmd', '.scr', '.pif', '.vbs', '.js']
            if any(filename.endswith(ext) for ext in dangerous_extensions):
                results['safe'] = False
                results['threats_found'].append('Potentially dangerous file extension')
            
            # Check file size anomalies
            if file.size > 100 * 1024 * 1024:  # 100MB
                results['recommendations'].append('Large file size - verify contents')
            
            # Check for embedded macros (basic check)
            if filename.endswith(('.docm', '.xlsm', '.pptm')):
                results['recommendations'].append('File contains macros - verify source')
            
        except Exception as e:
            logger.error(f"Security scan failed: {e}")
            results['safe'] = False
            results['threats_found'].append(f'Scan error: {str(e)}')
        
        return results